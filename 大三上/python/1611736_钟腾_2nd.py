'''
Created on 2018年10月18日

@author: zhongteng
'''
from docx import Document
from collections import OrderedDict
from pyexcel_io.io import save_data

def not_empty(s):
    '''去空串'''
    while '' in s:
        s.remove('')
    return s

def read_answer():
    '''读取参考答案内容'''
    
    document = Document("/Users/zhongteng/python/参考答案.docx") #读取word内容
    anslist = [] #设置答案列表
    
    ansstr = document.paragraphs[5].text.split() #读取第五行答案
    for j in ansstr:
        if j.isalpha():
            anslist.append(j.replace(',','').replace(' ',''))
            
    for i in range(8,12): #读取第8-12行答案
        ansstr = document.paragraphs[i].text.split('\t')
        not_empty(ansstr)
        for j in ansstr:
            res = j.split('.')
            for k in res:
                if not k.isalnum():
                    anslist.append(k.strip().replace(',','').replace(' ',''))
                    
    ansstr = document.paragraphs[17].text.split() #读取第17行答案
    for j in ansstr:
        if j.isalpha():
            anslist.append(j.replace(',','').replace(' ',''))
            
    ansstr = document.paragraphs[20].text.split() #读取第20行答案
    for j in ansstr:
        if j.isalpha():
            anslist.append(j.replace(',','').replace(' ',''))
            
    for i in range(23,25): #读取第23-25行答案
        ansstr = document.paragraphs[i].text.split('\t')
        not_empty(ansstr)
        for j in ansstr:
            res = j.split('.')
            for k in res:
                if k=='33':
                    res.remove('33')
                    anslist.append('.'.join(res))
                elif k=='nankai' or k=='edu' or k=='cn':
                    continue
                elif not k.isdigit():
                    anslist.append(k.strip().replace(',','').replace(' ',''))
                    
    ansstr = document.paragraphs[27].text.split('\t') #读取第27行答案
    not_empty(ansstr)
    for j in ansstr:
        res = j.split('.')
        for k in res:
            if not k.isdigit():
                anslist.append(k.strip().replace(',','').replace(' ' ,''))
                
    for i in range(30,32): #读取第30-32行答案
        ansstr = document.paragraphs[i].text.split('\t')
        not_empty(ansstr)
        for j in ansstr:
            res = j.split('.')
            for k in res:
                if not k.strip().isdigit():
                    anslist.append(k.strip().replace(',','').replace(' ',''))
                    
    ansstr = document.paragraphs[34].text.split() #读取第34行答案
    for j in ansstr:
        list1 = j.split('.')
        for k in list1:
            if k.isalpha():
                anslist.append(k.replace(',','').replace(' ',''))
    anslist = [string.upper() for string in anslist] #全部大写 
    return anslist
    
def read_paper():
    '''读取试卷'''
    
    document = Document("/Users/zhongteng/python/答题卡/5-赵.docx") #读取试卷内容
    tables = document.tables #获取word中的所有表格
    tablelist = []
    
    for i in range(1,3): #读取第2-3个表格
        for j in range(1,4):
            if j%2!=0:
                for k in range(10):
                    tablelist.append(tables[i].cell(j,k).text.replace(',','').replace(' ',''))
                    
    for i in range(3,5): #读取第3-4个表格
        for j in range(5):
            tablelist.append(tables[i].cell(j,1).text.replace(',','').replace(' ',''))
    tablelist = [string.upper() for string in tablelist] #全部大写
    return tablelist
 
def count_score(tablelist,anslist):
    '''比较试卷与答案，计算分数'''
    
    print(tablelist)
    print(anslist)
    total_score = 0
    parts_score = [0]*4
    
    for i in range(0,20): #第一大题分数
        if tablelist[i]=='':
            continue
        if tablelist[i] in anslist[i]:
            parts_score[0]+=1
    for i in range(20,45): #第二大题分数
        if tablelist[i]=='':
            continue
        if tablelist[i] == anslist[i]:
            print(tablelist[i])
            parts_score[1]+=2
    for i in range(45,50): #第三大题分数
        if tablelist[i]=='':
            continue
        if tablelist[i] in anslist[i]:
            parts_score[2]+=2
            
    total_score=sum(parts_score) #总分
    for i in range(0,4):
        print(parts_score[i])
    print(total_score)
    
    document = Document("/Users/zhongteng/python/答题卡/5-赵.docx") #读取试卷内容
    tables = document.tables #获取word中的所有表格
    tables[0].cell(1,1).text = str(parts_score[0])
    tables[0].cell(1,2).text = str(parts_score[1])
    tables[0].cell(1,3).text = str(parts_score[2])
    tables[0].cell(1,4).text = str(parts_score[3])
    tables[0].cell(1,5).text = str(total_score)
    document.save("/Users/zhongteng/python/答题卡/5-赵.docx")
    
def save_xls_file():
    '''写excel表单'''
    
    data = OrderedDict()
     
    sheet_1 = []
    row_1_data = [u"ID",u"昵称",u"等级"]
    row_2_data = [4,5,6]
     
    sheet_1.append(row_1_data) 
    sheet_1.append(row_2_data)
     
    data.update({u"这是xx表":sheet_1})  
     
    save_data("/Users/zhongteng/test.xls",data)
     
 
count_score(read_paper(), read_answer())